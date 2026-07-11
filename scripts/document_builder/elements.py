"""
Reusable document elements.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import (
    PRIMARY,
    SECONDARY,
    TEXT,
)


# ==========================================================
# Contact Card
# ==========================================================

def add_contact_card(
    document,
    location,
    phone,
    email,
    github,
    linkedin,
):
    """
    Add a professional contact section.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        f"📍 {location}\n"
        f"☎ {phone}\n"
        f"✉ {email}\n"
        f"GitHub: {github}\n"
        f"LinkedIn: {linkedin}"
    )

    run.font.name = "Calibri"
    run.font.size = 11
    run.font.color.rgb = TEXT

    document.add_paragraph()


# ==========================================================
# Signature Block
# ==========================================================

def add_signature(
    document,
    name,
    title,
):
    """
    Add a professional signature.
    """

    document.add_paragraph()

    paragraph = document.add_paragraph()

    paragraph.add_run("Kind regards,\n\n")

    name_run = paragraph.add_run(name)

    name_run.bold = True

    paragraph.add_run(f"\n{title}")


# ==========================================================
# Information Box
# ==========================================================

def add_information_box(
    document,
    heading,
    content,
):
    """
    Add a highlighted information box.
    """

    paragraph = document.add_paragraph()

    title = paragraph.add_run(
        heading + "\n"
    )

    title.bold = True

    title.font.color.rgb = PRIMARY

    body = paragraph.add_run(content)

    body.font.color.rgb = TEXT

    document.add_paragraph()


# ==========================================================
# Quote
# ==========================================================

def add_quote(
    document,
    text,
):
    """
    Add a quotation.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        f'"{text}"'
    )

    run.italic = True

    run.font.color.rgb = SECONDARY

    document.add_paragraph()


# ==========================================================
# Two Column Helper (Placeholder)
# ==========================================================

def add_two_column(
    document,
    left,
    right,
):
    """
    Temporary two-column layout.

    This will later become a proper table-based layout.
    """

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.style = "Table Grid"

    table.cell(
        0,
        0,
    ).text = left

    table.cell(
        0,
        1,
    ).text = right


# ==========================================================
# Skills Table
# ==========================================================

def add_skills_table(
    document,
    rows,
):
    """
    Create a two-column skills table.

    rows should be:

    [
        ("Programming", "Python, Git"),
        ("GIS", "ArcGIS Pro"),
    ]
    """

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.style = "Light Grid Accent 1"

    header = table.rows[0].cells

    header[0].text = "Category"

    header[1].text = "Skills"

    for category, skills in rows:

        cells = table.add_row().cells

        cells[0].text = category

        cells[1].text = skills

    document.add_paragraph()