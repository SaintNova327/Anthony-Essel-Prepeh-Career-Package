"""
Document layout utilities.

Responsible for:

- Page configuration
- Cover pages
- Headers
- Footers
- Page breaks
"""

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .styling import (
    PRIMARY,
    SECONDARY,
)


# ==========================================================
# Page Configuration
# ==========================================================

def configure_page(document):
    """
    Configure page margins and spacing.
    """

    section = document.sections[0]

    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    return section


# ==========================================================
# Cover Page
# ==========================================================

def add_cover_page(
    document,
    title,
    subtitle,
):
    """
    Add a professional cover page.
    """

    title_paragraph = document.add_paragraph()

    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title_paragraph.add_run(title)

    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    document.add_paragraph()

    subtitle_paragraph = document.add_paragraph()

    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle_paragraph.add_run(subtitle)

    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = SECONDARY

    document.add_page_break()


# ==========================================================
# Header
# ==========================================================

def add_header(
    document,
    text,
):
    """
    Add a simple header.
    """

    section = document.sections[0]

    header = section.header

    paragraph = header.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)

    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)


# ==========================================================
# Footer
# ==========================================================

def add_footer(
    document,
    text,
):
    """
    Add footer text.
    """

    section = document.sections[0]

    footer = section.footer

    paragraph = footer.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)

    run.font.name = "Calibri"
    run.font.size = Pt(9)


# ==========================================================
# Page Break
# ==========================================================

def add_page_break(document):
    """
    Insert a page break.
    """

    document.add_page_break()


# ==========================================================
# New Section
# ==========================================================

def add_section(document):
    """
    Start a new document section.
    """

    return document.add_section(
        WD_SECTION.NEW_PAGE
    )