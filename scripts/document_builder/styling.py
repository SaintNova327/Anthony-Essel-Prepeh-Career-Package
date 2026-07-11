"""
Professional styling utilities for DOCX documents.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


# ==========================================================
# Brand Colours
# ==========================================================

PRIMARY = RGBColor(10, 61, 98)      # Navy Blue
SECONDARY = RGBColor(184, 134, 11)  # Gold
TEXT = RGBColor(40, 40, 40)


# ==========================================================
# Create Document
# ==========================================================

def create_document(title):
    """
    Create a professionally styled Word document.
    """

    document = Document()

    # ------------------------------------------------------
    # Document Properties
    # ------------------------------------------------------

    core = document.core_properties

    core.author = "Anthony Essel Prepeh"
    core.title = title
    core.subject = "Professional Career Document"
    core.company = "Anthony Essel Prepeh Career Operating System"
    core.comments = (
        "Automatically generated using the Career Operating System."
    )
    core.language = "en-GH"

    # ------------------------------------------------------
    # Page Margins
    # ------------------------------------------------------

    section = document.sections[0]

    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # ------------------------------------------------------
    # Title
    # ------------------------------------------------------

    title_paragraph = document.add_paragraph()

    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title_paragraph.add_run(title)

    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = PRIMARY

    title_run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    # ------------------------------------------------------
    # Subtitle
    # ------------------------------------------------------

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "Geological Engineering\n"
        "Mining Technology\n"
        "Artificial Intelligence"
    )

    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = SECONDARY

    subtitle_run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    document.add_paragraph()
    document.add_paragraph()

    return document


# ==========================================================
# Heading
# ==========================================================

def add_heading(document, text, level=1):
    """
    Add a section heading.
    """

    heading = document.add_heading(level=level)

    run = heading.add_run(text)

    run.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY

    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    return heading


# ==========================================================
# Paragraph
# ==========================================================

def add_paragraph(document, text):
    """
    Add a formatted paragraph.
    """

    paragraph = document.add_paragraph()

    paragraph.space_after = Pt(8)
    paragraph.line_spacing = 1.15

    run = paragraph.add_run(text)

    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    return paragraph


# ==========================================================
# Bullet List
# ==========================================================

def add_bullet(document, text):
    """
    Add a bullet point.
    """

    paragraph = document.add_paragraph(
        style="List Bullet"
    )

    paragraph.space_after = Pt(3)

    run = paragraph.add_run(text)

    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    return paragraph


# ==========================================================
# Divider
# ==========================================================

def add_divider(document):
    """
    Add a decorative divider.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "────────────────────────────────────────"
    )

    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = SECONDARY

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Calibri",
    )

    return paragraph